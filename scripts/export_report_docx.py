#!/usr/bin/env python3
import argparse
import json
import os
import re
import sys
import urllib.parse
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(15, 118, 110)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(100, 116, 139)
LINE = RGBColor(226, 232, 240)


def require_env(name):
    value = os.environ.get(name)
    if not value:
        raise RuntimeError(f"{name} is required")
    return value


def supabase_get(path):
    base_url = require_env("SUPABASE_URL").rstrip("/")
    key = require_env("SUPABASE_SERVICE_ROLE_KEY")
    request = urllib.request.Request(
        f"{base_url}/rest/v1/{path}",
        headers={
            "apikey": key,
            "Authorization": f"Bearer {key}",
            "Accept": "application/json",
        },
    )
    with urllib.request.urlopen(request, timeout=30) as response:
        return json.loads(response.read().decode("utf-8"))


def latest_report():
    params = urllib.parse.urlencode(
        {
            "select": "id,title,report_type,report_date,markdown_body,created_at",
            "report_type": "eq.daily",
            "status": "eq.ready",
            "order": "report_date.desc,created_at.desc",
            "limit": "20",
        }
    )
    rows = supabase_get(f"reports?{params}")
    for row in rows:
        if re.match(r"^\d{4}-\d{2}-\d{2}$", str(row.get("report_date", ""))):
            return row
    raise RuntimeError("No public daily report found")


def report_by_id(report_id):
    params = urllib.parse.urlencode(
        {
            "select": "id,title,report_type,report_date,markdown_body,created_at",
            "id": f"eq.{report_id}",
            "status": "eq.ready",
            "limit": "1",
        }
    )
    rows = supabase_get(f"reports?{params}")
    if not rows:
        raise RuntimeError(f"Report not found: {report_id}")
    return rows[0]


def report_by_date(report_date):
    params = urllib.parse.urlencode(
        {
            "select": "id,title,report_type,report_date,markdown_body,created_at",
            "report_type": "eq.daily",
            "report_date": f"eq.{report_date}",
            "status": "eq.ready",
            "order": "created_at.desc",
            "limit": "1",
        }
    )
    rows = supabase_get(f"reports?{params}")
    if not rows:
        raise RuntimeError(f"Report not found for date: {report_date}")
    return rows[0]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_border(cell, color="E2E8F0"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_paragraph_border(paragraph, color="14B8A6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), color)
    left.set(qn("w:space"), "8")
    borders.append(left)
    p_pr.append(borders)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0F766E")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(fonts)
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def normalize_links(text):
    text = re.sub(r"(https?:)\s*/\s*/", r"\1//", text)
    return re.sub(
        r"\]\((https?://[^)]*?)\)",
        lambda match: "](" + re.sub(r"\s+", "", match.group(1)) + ")",
        text,
        flags=re.S,
    )


def clean_inline(text):
    return re.sub(r"`([^`]+)`", r"\1", re.sub(r"\*\*([^*]+)\*\*", r"\1", text))


def add_inline(paragraph, text, size=10.5, bold=False, color=INK):
    text = normalize_links(clean_inline(text))
    pos = 0
    for match in re.finditer(r"\[([^\]]+)\]\((https?://[^)\s]+)\)", text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = color
        add_hyperlink(paragraph, match.group(1), match.group(2))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color


def add_paragraph(doc, text="", size=10.5, bold=False, color=INK, before=0, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.35
    add_inline(paragraph, text, size=size, bold=bold, color=color)
    return paragraph


def apply_base_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = INK


def add_title_block(doc, report):
    title = report.get("title") or "MornInvest 美股科技晨报"
    date = report.get("report_date") or ""
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("MORNINVEST")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    add_inline(paragraph, title, size=22, bold=True)

    add_paragraph(doc, "基于公开信息整理，不构成投资建议。", size=9.5, color=MUTED, after=14)


def add_table(doc, table_lines):
    rows = []
    for line in table_lines:
        if re.match(r"^\|\s*-+", line):
            continue
        cells = [clean_inline(cell.strip()) for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            if row_index == 0:
                set_cell_shading(cell, "F1F5F9")
            text = row[col_index] if col_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(
                paragraph,
                text,
                size=8.8,
                bold=row_index == 0,
                color=INK if row_index == 0 else INK,
            )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def render_markdown(doc, markdown):
    lines = normalize_links(markdown).replace("\r\n", "\n").split("\n")
    table_lines = []
    started = False

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_table(doc, table_lines)
            table_lines = []

    for raw_line in lines:
        line = raw_line.strip()
        if not started:
            if re.match(r"^##\s+0\.\s+", line):
                started = True
            else:
                continue

        if not line:
            flush_table()
            continue
        if re.match(r"^-{3,}$", line):
            flush_table()
            continue
        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            continue

        flush_table()

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1:
                continue
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(12 if level == 2 else 8)
            paragraph.paragraph_format.space_after = Pt(6)
            if level == 3:
                set_paragraph_border(paragraph)
            add_inline(paragraph, text, size=15 if level == 2 else 12.5, bold=True)
            continue

        bullet = re.match(r"^[-*+]\s+(.+)$", line)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
            add_inline(paragraph, bullet.group(1), size=10)
            continue

        ordered = re.match(r"^\d+\.\s+(.+)$", line)
        if ordered:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(3)
            add_inline(paragraph, ordered.group(1), size=10)
            continue

        field = re.match(r"^([\u4e00-\u9fffA-Za-z0-9 /-]{2,32}[：:])\s*(.*)$", line)
        if field:
            paragraph = add_paragraph(doc, "", size=10, before=2, after=2)
            add_inline(paragraph, field.group(1), size=10, bold=True, color=MUTED)
            if field.group(2):
                add_inline(paragraph, " " + field.group(2), size=10.5, color=INK)
            continue

        add_paragraph(doc, line, size=10.5, after=5)

    flush_table()


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(8)
    run = footer.add_run("MornInvest · 基于公开信息整理，不构成投资建议")
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def safe_filename(value):
    value = re.sub(r"[^\w\u4e00-\u9fff.-]+", "-", value, flags=re.U).strip("-")
    return value or "morninvest-report"


def build_docx(report, output_dir):
    doc = Document()
    apply_base_styles(doc)
    add_title_block(doc, report)
    render_markdown(doc, report.get("markdown_body") or "")
    add_footer(doc)

    output_dir.mkdir(parents=True, exist_ok=True)
    filename = safe_filename(f"morninvest-{report.get('report_date', 'latest')}-雪球版.docx")
    output_path = output_dir / filename
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Export the latest MornInvest report as a DOCX file.")
    parser.add_argument("--report-id", default=os.environ.get("REPORT_ID"), help="Optional Supabase report id")
    parser.add_argument("--report-date", default=os.environ.get("REPORT_DATE"), help="Optional YYYY-MM-DD report date")
    parser.add_argument("--out-dir", default="report-output", help="Output directory")
    args = parser.parse_args()

    if args.report_id:
        report = report_by_id(args.report_id)
    elif args.report_date:
        report = report_by_date(args.report_date)
    else:
        report = latest_report()
    output_path = build_docx(report, Path(args.out_dir))
    print(json.dumps({"report_id": report["id"], "report_date": report.get("report_date"), "path": str(output_path)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    try:
        main()
    except Exception as error:
        print(str(error), file=sys.stderr)
        sys.exit(1)
